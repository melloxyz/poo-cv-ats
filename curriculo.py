import PyPDF2
import docx
from typing import Optional, Dict, Any
import tempfile
import os
import re

class Curriculo:
    """Classe responsável pela extração dos dados do currículo"""
    
    def __init__(self, arquivo_upload):
        self.arquivo = arquivo_upload
        self.nome_arquivo = arquivo_upload.name if arquivo_upload else None
        self.tipo_arquivo = self._identificar_tipo_arquivo()
        self.texto_extraido = None
        self.dados_estruturados = None
        self.metadados = {}
        
        # Inicializa o extrator inteligente se disponível
        try:
            from extrator_inteligente import ExtratorInteligente
            self.extrator_ia = ExtratorInteligente()
            self.usar_ia = True
        except ImportError:
            print("⚠️ Extrator inteligente não disponível, usando extração básica")
            self.extrator_ia = None
            self.usar_ia = False
    
    def extrair_texto(self) -> Dict[str, Any]:
        """Extrai texto e dados estruturados do currículo"""
        if not self.arquivo:
            return {
                "sucesso": False,
                "erro": "Nenhum arquivo foi carregado"
            }
        
        try:
            # 1. Extração de texto básica
            if self.tipo_arquivo == 'pdf':
                self.texto_extraido = self._extrair_texto_pdf()
            elif self.tipo_arquivo == 'docx':
                self.texto_extraido = self._extrair_texto_docx()
            else:
                return {
                    "sucesso": False,
                    "erro": f"Tipo de arquivo não suportado: {self.tipo_arquivo}"
                }
            
            # 2. Valida se conseguiu extrair texto
            if not self.texto_extraido or len(self.texto_extraido.strip()) < 10:
                return {
                    "sucesso": False,
                    "erro": "Não foi possível extrair texto válido do arquivo"
                }
            
            # 3. Extração de metadados básicos
            self._extrair_metadados()
            
            # 4. Extração inteligente com IA (se disponível)
            if self.usar_ia and self.extrator_ia:
                try:
                    print("🤖 Iniciando extração inteligente com IA...")
                    self.dados_estruturados = self.extrator_ia.extrair_dados_completos(self.texto_extraido)
                    print("✅ Extração inteligente concluída com sucesso!")
                except Exception as e:
                    print(f"⚠️ Falha na extração inteligente: {e}")
                    print("🔄 Continuando com extração básica...")
                    self.dados_estruturados = self._extrair_dados_basicos_regex()
            else:
                # Fallback: extração básica com regex
                self.dados_estruturados = self._extrair_dados_basicos_regex()
            
            return {
                "sucesso": True,
                "texto": self.texto_extraido,
                "dados_estruturados": self.dados_estruturados,
                "metadados": self.metadados,
                "metodo_extracao": "IA" if self.usar_ia else "REGEX"
            }
            
        except Exception as e:
            return {
                "sucesso": False,
                "erro": f"Erro ao extrair dados do {self.tipo_arquivo.upper()}: {str(e)}"
            }
    
    def _extrair_dados_basicos_regex(self) -> Dict[str, Any]:
        """
        Extração básica usando regex como fallback da IA.
        
        Returns:
            Dict: Dados básicos extraídos com regex
        """
        if not self.texto_extraido:
            return {"qualidade_extracao": "FALHA", "erro": "Sem texto para processar"}
        
        texto = self.texto_extraido
        
        # Extração de dados básicos com regex melhorado
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        telefone_pattern = r'(\(?\d{2}\)?[\s-]?\d{4,5}[\s-]?\d{4})'
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/profile/)([A-Za-z0-9-_]+)'
        github_pattern = r'(?:github\.com/)([A-Za-z0-9-_]+)'
        
        emails = re.findall(email_pattern, texto)
        telefones = re.findall(telefone_pattern, texto)
        linkedin = re.findall(linkedin_pattern, texto, re.IGNORECASE)
        github = re.findall(github_pattern, texto, re.IGNORECASE)
        
        # Extração de nome (melhorada)
        nome = self._extrair_nome_candidato(texto)
        
        # Extração de habilidades técnicas básicas
        tecnologias = self._extrair_tecnologias_basicas(texto)
        
        # Estimativa de experiência
        experiencia_anos = self._estimar_experiencia(texto)
        
        return {
            "dados_pessoais": {
                "nome_completo": nome,
                "email": emails[0] if emails else "Não identificado",
                "telefone": telefones[0] if telefones else "Não identificado",
                "linkedin": f"https://linkedin.com/in/{linkedin[0]}" if linkedin else "Não identificado",
                "github": f"https://github.com/{github[0]}" if github else "Não identificado",
                "endereco": "Não identificado"
            },
            "habilidades_competencias": {
                "linguagens_programacao": tecnologias["linguagens"],
                "frameworks_bibliotecas": tecnologias["frameworks"],
                "ferramentas": tecnologias["ferramentas"],
                "nivel_tecnico_geral": self._classificar_nivel_tecnico(experiencia_anos, tecnologias)
            },
            "experiencia_profissional": {
                "tempo_total_experiencia_anos": experiencia_anos,
                "experiencias": [],
                "areas_de_atuacao": self._identificar_areas_atuacao(texto)
            },
            "qualidade_extracao": "REGEX_BASICA",
            "timestamp_extracao": "2025-09-10",
            "metricas_calculadas": {
                "score_completude": self._calcular_completude_basica(nome, emails, tecnologias),
                "nivel_senioridade_calculado": self._classificar_senioridade(experiencia_anos)
            }
        }
    
    def _extrair_nome_candidato(self, texto: str) -> str:
        """
        Extrai o nome do candidato usando heurísticas melhoradas.
        """
        linhas = texto.split('\n')
        
        for linha in linhas[:10]:  # Verifica primeiras 10 linhas
            linha_limpa = linha.strip()
            
            # Ignora linhas muito curtas ou com padrões de não-nome
            if (linha_limpa and 
                len(linha_limpa) > 5 and 
                len(linha_limpa) < 50 and
                not re.search(r'@|tel|cel|phone|email|cv|curriculum|endereço|www|http', linha_limpa.lower())):
                
                # Verifica se não tem muitos números (evita datas, telefones)
                numeros = len(re.findall(r'\d', linha_limpa))
                if numeros / len(linha_limpa) < 0.3:  # Menos de 30% números
                    
                    # Verifica se parece com nome (tem pelo menos 2 palavras)
                    palavras = linha_limpa.split()
                    if len(palavras) >= 2:
                        return linha_limpa
        
        return "Não identificado"
    
    def _extrair_tecnologias_basicas(self, texto: str) -> Dict[str, list]:
        """
        Extrai tecnologias mencionadas usando lista pré-definida.
        """
        texto_lower = texto.lower()
        
        linguagens = [
            'python', 'javascript', 'java', 'php', 'c++', 'c#', 'ruby', 'go', 
            'typescript', 'kotlin', 'swift', 'scala', 'rust', 'r', 'matlab'
        ]
        
        frameworks = [
            'react', 'angular', 'vue', 'django', 'flask', 'spring', 'laravel',
            'express', 'next.js', 'nuxt', 'bootstrap', 'tailwind', 'fastapi'
        ]
        
        ferramentas = [
            'git', 'docker', 'kubernetes', 'jenkins', 'aws', 'azure', 'gcp',
            'mongodb', 'postgresql', 'mysql', 'redis', 'elasticsearch', 'nginx'
        ]
        
        linguagens_encontradas = [tech for tech in linguagens if tech in texto_lower]
        frameworks_encontrados = [tech for tech in frameworks if tech in texto_lower]
        ferramentas_encontradas = [tech for tech in ferramentas if tech in texto_lower]
        
        return {
            "linguagens": linguagens_encontradas,
            "frameworks": frameworks_encontrados,
            "ferramentas": ferramentas_encontradas
        }
    
    def _estimar_experiencia(self, texto: str) -> float:
        """
        Estima anos de experiência baseado em padrões no texto.
        """
        # Procura por menções diretas de experiência
        padroes_experiencia = [
            r'(\d+)\s*anos?\s*de\s*experiência',
            r'experiência\s*de\s*(\d+)\s*anos?',
            r'(\d+)\s*years?\s*of\s*experience',
            r'experience\s*of\s*(\d+)\s*years?'
        ]
        
        anos_encontrados = []
        for padrao in padroes_experiencia:
            matches = re.findall(padrao, texto.lower())
            anos_encontrados.extend([int(match) for match in matches])
        
        if anos_encontrados:
            return max(anos_encontrados)  # Pega o maior valor encontrado
        
        # Estimativa baseada em número de empregos mencionados
        empresas_patterns = r'(empresa|company|corp|ltd|inc|sa|ltda)'
        empresas_count = len(re.findall(empresas_patterns, texto.lower()))
        
        if empresas_count > 3:
            return 5.0  # Assume 5+ anos se menciona muitas empresas
        elif empresas_count > 1:
            return 3.0  # Assume ~3 anos
        else:
            return 1.0  # Assume iniciante
    
    def _identificar_areas_atuacao(self, texto: str) -> list:
        """
        Identifica áreas de atuação baseado em palavras-chave.
        """
        texto_lower = texto.lower()
        areas = []
        
        areas_mapping = {
            'desenvolvimento': ['desenvolvimento', 'programação', 'software', 'developer'],
            'frontend': ['frontend', 'front-end', 'html', 'css', 'javascript', 'react'],
            'backend': ['backend', 'back-end', 'server', 'api', 'database'],
            'fullstack': ['fullstack', 'full-stack', 'full stack'],
            'devops': ['devops', 'deployment', 'docker', 'kubernetes', 'ci/cd'],
            'data_science': ['data science', 'machine learning', 'analytics', 'data analyst'],
            'mobile': ['mobile', 'android', 'ios', 'react native', 'flutter'],
            'gestao': ['gerente', 'manager', 'liderança', 'coordenador', 'supervisor']
        }
        
        for area, keywords in areas_mapping.items():
            if any(keyword in texto_lower for keyword in keywords):
                areas.append(area.replace('_', ' ').title())
        
        return areas[:5]  # Limita a 5 áreas
    
    def _classificar_nivel_tecnico(self, anos_exp: float, tecnologias: Dict) -> str:
        """
        Classifica nível técnico baseado em experiência e tecnologias conhecidas.
        """
        total_techs = (len(tecnologias["linguagens"]) + 
                      len(tecnologias["frameworks"]) + 
                      len(tecnologias["ferramentas"]))
        
        if anos_exp >= 5 and total_techs >= 10:
            return "Senior"
        elif anos_exp >= 2 and total_techs >= 5:
            return "Pleno"
        else:
            return "Junior"
    
    def _calcular_completude_basica(self, nome: str, emails: list, tecnologias: Dict) -> int:
        """
        Calcula score básico de completude dos dados (0-100).
        """
        pontos = 0
        
        if nome != "Não identificado":
            pontos += 25
        if emails:
            pontos += 25
        if tecnologias["linguagens"]:
            pontos += 25
        if len(tecnologias["linguagens"]) > 2:
            pontos += 25
        
        return min(100, pontos)
    
    def _classificar_senioridade(self, anos_exp: float) -> str:
        """
        Classifica senioridade baseado em anos de experiência.
        """
        if anos_exp < 2:
            return "Junior"
        elif anos_exp < 5:
            return "Pleno"
        elif anos_exp < 8:
            return "Senior"
        else:
            return "Especialista"
    
    def _identificar_tipo_arquivo(self) -> Optional[str]:
        """
        Identifica o tipo do arquivo baseado na extensão.
        
        Returns:
            str: Tipo do arquivo ('pdf' ou 'docx')
        """
        if not self.nome_arquivo:
            return None
            
        extensao = self.nome_arquivo.lower().split('.')[-1]
        
        if extensao == 'pdf':
            return 'pdf'
        elif extensao in ['docx', 'doc']:
            return 'docx'
        else:
            return None
    
    def _extrair_texto_pdf(self) -> str:
        """
        Extrai texto de arquivo PDF usando PyPDF2.
        
        Returns:
            str: Texto extraído do PDF
        """
        texto_completo = ""
        
        # Salva o arquivo temporariamente para leitura
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(self.arquivo.getvalue())
            temp_path = temp_file.name
        
        try:
            with open(temp_path, 'rb') as arquivo_pdf:
                reader = PyPDF2.PdfReader(arquivo_pdf)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        texto_pagina = page.extract_text()
                        if texto_pagina.strip():
                            texto_completo += texto_pagina + "\n"
                    except Exception as e:
                        print(f"Erro ao extrair texto da página {page_num + 1}: {e}")
                        continue
        
        finally:
            # Remove arquivo temporário
            if os.path.exists(temp_path):
                os.unlink(temp_path)
        
        return texto_completo.strip()
    
    def _extrair_texto_docx(self) -> str:
        """
        Extrai texto de arquivo DOCX usando python-docx.
        
        Returns:
            str: Texto extraído do DOCX
        """
        texto_completo = ""
        
        # Salva o arquivo temporariamente para leitura
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(self.arquivo.getvalue())
            temp_path = temp_file.name
        
        try:
            doc = docx.Document(temp_path)
            
            # Extrai texto de todos os parágrafos
            for paragrafo in doc.paragraphs:
                if paragrafo.text.strip():
                    texto_completo += paragrafo.text + "\n"
            
            # Extrai texto de tabelas se existirem
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        if celula.text.strip():
                            texto_completo += celula.text + "\n"
        
        finally:
            # Remove arquivo temporário
            if os.path.exists(temp_path):
                os.unlink(temp_path)
        
        return texto_completo.strip()
    
    def _extrair_metadados(self):
        """
        Extrai metadados básicos do currículo extraído.
        """
        if not self.texto_extraido:
            return
        
        import re
        
        # Buscar email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, self.texto_extraido)
        
        # Buscar telefone
        telefone_pattern = r'(\(?\d{2}\)?[\s-]?\d{4,5}[\s-]?\d{4})|(\+?\d{2}[\s-]?\(?\d{2}\)?[\s-]?\d{4,5}[\s-]?\d{4})'
        telefones = re.findall(telefone_pattern, self.texto_extraido)
        
        # Buscar possível nome (primeira linha que não é email/telefone)
        linhas = self.texto_extraido.split('\n')
        nome_candidato = "Não identificado"
        for linha in linhas[:5]:  # Verifica as primeiras 5 linhas
            linha = linha.strip()
            if linha and len(linha) > 5 and not re.search(r'@|tel|cel|phone|email', linha.lower()):
                if not any(char.isdigit() for char in linha[:10]):  # Se não tem números no início
                    nome_candidato = linha
                    break
        
        self.metadados = {
            'tamanho_texto': len(self.texto_extraido),
            'numero_linhas': len(self.texto_extraido.split('\n')),
            'numero_palavras': len(self.texto_extraido.split()),
            'emails_encontrados': emails[:3] if emails else [],
            'telefones_encontrados': [t[0] or t[1] for t in telefones[:2]] if telefones else [],
            'nome_candidato': nome_candidato,
            'palavras_chave_tech': self._extrair_palavras_chave_tecnicas(),
            'tipo_arquivo': self.tipo_arquivo,
            'nome_arquivo': self.nome_arquivo
        }
    
    def _extrair_palavras_chave_tecnicas(self):
        """
        Extrai palavras-chave técnicas comuns do currículo.
        
        Returns:
            list: Lista de tecnologias/habilidades encontradas
        """
        if not self.texto_extraido:
            return []
        
        # Tecnologias e habilidades comuns
        tecnologias = [
            'python', 'javascript', 'java', 'c++', 'c#', 'php', 'ruby', 'go', 'rust',
            'html', 'css', 'react', 'angular', 'vue', 'node', 'express', 'django', 
            'flask', 'spring', 'laravel', 'sql', 'mysql', 'postgresql', 'mongodb',
            'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'git', 'github', 'gitlab',
            'linux', 'windows', 'mac', 'android', 'ios', 'flutter', 'react native',
            'machine learning', 'ai', 'data science', 'excel', 'power bi', 'tableau'
        ]
        
        texto_lower = self.texto_extraido.lower()
        encontradas = []
        
        for tech in tecnologias:
            if tech in texto_lower:
                encontradas.append(tech.title())
        
        return encontradas[:10]  # Retorna no máximo 10
    
    def obter_resumo(self) -> Dict[str, Any]:
        """
        Retorna um resumo dos dados extraídos do currículo.
        
        Returns:
            Dict: Resumo com metadados e preview do texto
        """
        if not self.texto_extraido:
            return {"erro": "Texto não foi extraído ainda"}
        
        # Preview do texto (primeiras 200 palavras)
        palavras = self.texto_extraido.split()
        preview = ' '.join(palavras[:200])
        if len(palavras) > 200:
            preview += "..."
        
        return {
            "metadados": self.metadados,
            "preview_texto": preview,
            "status": "Texto extraído com sucesso"
        }
    
    def validar_arquivo(self) -> Dict[str, Any]:
        """
        Valida se o arquivo carregado é válido para processamento.
        
        Returns:
            Dict: Status de validação
        """
        if not self.arquivo:
            return {"valido": False, "erro": "Nenhum arquivo carregado"}
        
        if not self.tipo_arquivo:
            return {"valido": False, "erro": f"Tipo de arquivo não suportado: {self.nome_arquivo}"}
        
        # Verifica tamanho do arquivo (max 10MB)
        tamanho_mb = len(self.arquivo.getvalue()) / (1024 * 1024)
        if tamanho_mb > 10:
            return {"valido": False, "erro": f"Arquivo muito grande: {tamanho_mb:.1f}MB (máximo: 10MB)"}
        
        return {
            "valido": True, 
            "tipo": self.tipo_arquivo,
            "tamanho_mb": round(tamanho_mb, 2),
            "nome": self.nome_arquivo
        }